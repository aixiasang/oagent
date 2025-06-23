import os
import json
import csv
from pathlib import Path
from typing import List, Dict, Any, Optional, Union
import requests
from bs4 import BeautifulSoup
from .base_rag import Document
from .text_splitters import get_text_splitter, TextSplitter


class DocumentLoader:
    """增强的文档加载器，支持多种文档格式和智能分割策略"""
    def __init__(
        self,
        chunk_size: int = 1000,
        chunk_overlap: int = 200,
        splitter_type: str = "recursive"
    ):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.splitter_type = splitter_type
        self.text_splitter = get_text_splitter(
            splitter_type=splitter_type,
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap
        )
    
    def load_documents(
        self,
        sources: List[Union[str, Dict[str, Any]]]
    ) -> List[Document]:
        """批量加载文档"""
        documents = []
        
        for source in sources:
            if isinstance(source, str):
                if source.startswith(('http://', 'https://')):
                    docs = self.load_url_documents(source)
                else:
                    docs = self.load_file_documents(source)
            elif isinstance(source, dict):
                content = source.get('content', '')
                metadata = source.get('metadata', {})
                docs = self.create_documents_from_text(content, metadata)
            else:
                continue
            
            documents.extend(docs)
        
        return documents
    
    def load_file_documents(self, file_path: str, metadata: Optional[Dict] = None) -> List[Document]:
        """从文件加载文档"""
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        # 基础元数据
        base_metadata = {
            "source": str(file_path),
            "source_type": "file",
            "file_name": file_path.name,
            "file_extension": file_path.suffix.lower(),
            "file_size": file_path.stat().st_size,
        }
        
        if metadata:
            base_metadata.update(metadata)
        
        # 根据文件类型选择合适的分割器
        extension = file_path.suffix.lower()
        splitter = self._get_splitter_for_extension(extension)
        
        # 加载内容
        content = self._load_file_content(file_path)
        
        # 创建文档
        return self._create_documents_with_splitter(content, base_metadata, splitter)
    
    def load_url_documents(self, url: str, metadata: Optional[Dict] = None) -> List[Document]:
        """从URL加载文档"""
        content = self._load_url_content(url)
        
        base_metadata = {
            "source": url,
            "source_type": "url",
        }
        
        if metadata:
            base_metadata.update(metadata)
        
        # 对网页内容使用HTML分割器
        splitter = get_text_splitter(
            splitter_type="html",
            chunk_size=self.chunk_size,
            chunk_overlap=self.chunk_overlap
        )
        
        return self._create_documents_with_splitter(content, base_metadata, splitter)
    
    def create_documents_from_text(
        self,
        text: str,
        metadata: Optional[Dict] = None,
        splitter_type: Optional[str] = None
    ) -> List[Document]:
        """从文本创建文档"""
        base_metadata = {
            "source_type": "text",
        }
        
        if metadata:
            base_metadata.update(metadata)
        
        # 选择分割器
        if splitter_type:
            splitter = get_text_splitter(
                splitter_type=splitter_type,
                chunk_size=self.chunk_size,
                chunk_overlap=self.chunk_overlap
            )
        else:
            splitter = self.text_splitter
        
        return self._create_documents_with_splitter(text, base_metadata, splitter)
    
    def _get_splitter_for_extension(self, extension: str) -> TextSplitter:
        """根据文件扩展名获取合适的分割器"""
        splitter_map = {
            '.md': 'markdown',
            '.py': 'python',
            '.js': 'javascript',
            '.ts': 'javascript',
            '.html': 'html',
            '.htm': 'html',
            '.tex': 'latex',
            '.csv': 'recursive',  # CSV用递归分割器
            '.json': 'recursive',  # JSON用递归分割器
        }
        
        splitter_type = splitter_map.get(extension, 'recursive')
        
        return get_text_splitter(
            splitter_type=splitter_type,
            chunk_size=self.chunk_size,
            chunk_overlap=self.chunk_overlap
        )
    
    def _create_documents_with_splitter(
        self,
        content: str,
        metadata: Dict,
        splitter: TextSplitter
    ) -> List[Document]:
        """使用指定分割器创建文档"""
        chunks = splitter.split_text(content)
        documents = []
        
        for i, chunk in enumerate(chunks):
            chunk_metadata = metadata.copy()
            chunk_metadata.update({
                "chunk_index": i,
                "total_chunks": len(chunks),
                "chunk_size": len(chunk),
            })
            
            documents.append(Document(
                page_content=chunk,
                metadata=chunk_metadata
            ))
        
        return documents
    
    def _load_file_content(self, file_path: Path) -> str:
        """加载文件内容"""
        extension = file_path.suffix.lower()
        
        if extension == '.pdf':
            return self._load_pdf(file_path)
        elif extension in ['.doc', '.docx']:
            return self._load_word(file_path)
        elif extension == '.json':
            return self._load_json(file_path)
        elif extension == '.csv':
            return self._load_csv(file_path)
        elif extension in ['.txt', '.md', '.py', '.js', '.html', '.css', '.tex', '.rst']:
            return self._load_text(file_path)
        else:
            # 尝试作为文本文件加载
            return self._load_text(file_path)
    
    def _load_url_content(self, url: str) -> str:
        """加载网页内容"""
        try:
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
            }
            response = requests.get(url, headers=headers, timeout=30)
            response.raise_for_status()
            
            soup = BeautifulSoup(response.content, 'html.parser')
            
            # 移除脚本和样式标签
            for script in soup(["script", "style", "nav", "footer", "header"]):
                script.decompose()
            
            # 提取主要内容
            main_content = soup.find('main') or soup.find('article') or soup.find('div', class_=lambda x: x and 'content' in x.lower()) or soup.body
            
            if main_content:
                text = main_content.get_text()
            else:
                text = soup.get_text()
            
            # 清理文本
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = '\n'.join(chunk for chunk in chunks if chunk)
            
            return text
            
        except Exception as e:
            raise Exception(f"无法加载URL {url}: {str(e)}")
    
    def _load_pdf(self, file_path: Path) -> str:
        """加载PDF文件"""
        try:
            import PyPDF2
            
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                text = ""
                for page_num, page in enumerate(reader.pages):
                    page_text = page.extract_text()
                    if page_text:
                        text += f"\n--- 第{page_num + 1}页 ---\n{page_text}\n"
                return text
        except ImportError:
            raise ImportError("请安装PyPDF2: pip install PyPDF2")
        except Exception as e:
            try:
                # 尝试使用pdfplumber
                import pdfplumber
                
                with pdfplumber.open(file_path) as pdf:
                    text = ""
                    for page_num, page in enumerate(pdf.pages):
                        page_text = page.extract_text()
                        if page_text:
                            text += f"\n--- 第{page_num + 1}页 ---\n{page_text}\n"
                    return text
            except ImportError:
                raise Exception(f"PDF读取失败，请安装PyPDF2或pdfplumber: {str(e)}")
            except Exception as e2:
                raise Exception(f"PDF读取失败: {str(e2)}")
    
    def _load_word(self, file_path: Path) -> str:
        """加载Word文档"""
        try:
            import docx
            
            doc = docx.Document(file_path)
            text = ""
            
            # 提取段落
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # 提取表格
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        text += row_text + "\n"
            
            return text
        except ImportError:
            raise ImportError("请安装python-docx: pip install python-docx")
        except Exception as e:
            raise Exception(f"Word文档读取失败: {str(e)}")
    
    def _load_json(self, file_path: Path) -> str:
        """加载JSON文件"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                data = json.load(file)
                
                # 如果是对象数组，转换为更可读的格式
                if isinstance(data, list) and len(data) > 0 and isinstance(data[0], dict):
                    formatted_text = ""
                    for i, item in enumerate(data):
                        formatted_text += f"\n--- 记录 {i + 1} ---\n"
                        for key, value in item.items():
                            formatted_text += f"{key}: {value}\n"
                    return formatted_text
                else:
                    return json.dumps(data, ensure_ascii=False, indent=2)
        except Exception as e:
            raise Exception(f"JSON文件读取失败: {str(e)}")
    
    def _load_csv(self, file_path: Path) -> str:
        """加载CSV文件"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                reader = csv.DictReader(file)
                content = []
                
                # 添加表头
                if reader.fieldnames:
                    content.append(" | ".join(reader.fieldnames))
                    content.append("-" * 50)
                
                # 添加数据行
                for row_num, row in enumerate(reader):
                    row_text = " | ".join(str(value) for value in row.values())
                    content.append(f"行{row_num + 1}: {row_text}")
                
                return "\n".join(content)
        except Exception as e:
            # 尝试普通CSV读取
            try:
                with open(file_path, 'r', encoding='utf-8') as file:
                    reader = csv.reader(file)
                    content = []
                    for row_num, row in enumerate(reader):
                        row_text = " | ".join(row)
                        content.append(f"行{row_num + 1}: {row_text}")
                    return "\n".join(content)
            except Exception as e2:
                raise Exception(f"CSV文件读取失败: {str(e2)}")
    
    def _load_text(self, file_path: Path) -> str:
        """加载文本文件"""
        encodings = ['utf-8', 'gbk', 'gb2312', 'utf-16', 'latin-1']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    return file.read()
            except UnicodeDecodeError:
                continue
            except Exception as e:
                raise Exception(f"文本文件读取失败: {str(e)}")
        
        raise Exception(f"无法识别文件编码: {file_path}")
    
    def load_directory(
        self,
        directory_path: str,
        glob_pattern: str = "**/*",
        exclude_patterns: Optional[List[str]] = None,
        metadata: Optional[Dict] = None
    ) -> List[Document]:
        """批量加载目录中的文件"""
        directory = Path(directory_path)
        
        if not directory.exists():
            raise FileNotFoundError(f"目录不存在: {directory_path}")
        
        exclude_patterns = exclude_patterns or [
            "**/.git/**",
            "**/node_modules/**",
            "**/__pycache__/**",
            "**/*.pyc",
            "**/.DS_Store"
        ]
        
        all_documents = []
        
        for file_path in directory.glob(glob_pattern):
            if file_path.is_file():
                # 检查是否应该排除
                should_exclude = False
                for pattern in exclude_patterns:
                    if file_path.match(pattern):
                        should_exclude = True
                        break
                
                if should_exclude:
                    continue
                
                try:
                    file_metadata = metadata.copy() if metadata else {}
                    file_metadata["directory"] = str(directory)
                    
                    docs = self.load_file_documents(str(file_path), file_metadata)
                    all_documents.extend(docs)
                except Exception as e:
                    print(f"警告: 无法加载文件 {file_path}: {str(e)}")
                    continue
        
        return all_documents 